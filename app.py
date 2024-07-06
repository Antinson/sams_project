from flask import Flask, render_template, request, flash, send_file, redirect, url_for
import docx
import tempfile
import re

app = Flask(__name__)
app.secret_key = "testing123"

def my_file_new(filename):
    return send_file(filename, as_attachment=True)

@app.route('/', methods = ['GET', 'POST'])
def index():
    return render_template('index.html')


@app.route('/templates', methods=['GET', 'POST'])
def templates():
    new_template = request.form['temps']
    return redirect(url_for("main", template = new_template))

@app.route('/<template>', methods = ['GET', 'POST'])
def main(template):

    if str(template) != 'favicon.ico':

        template_name = str(template + '.docx')
        document_data = read_document(template_name)

        print(document_data)

        display_data = []

        for value in document_data:
            if value[1:len(value)-1] not in display_data:
                display_data.append(value[1:len(value)-1])

        for value in display_data:
            flash(value)
        
        if request.method == 'POST':
            new_dict = {}

            doc = docx.Document(template_name)
            new_document = str(request.form['document']) + '.docx'

            for value in document_data:
                new_dict[value] = str(request.form[value[1:len(value)-1]])
            

            for i in new_dict:
                for p in doc.paragraphs:
                    if p.text.find(i)>=0:
                        p.text=p.text.replace(i, new_dict[i])

            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                doc.save(tmp_file.name)
                tmp_file_path = tmp_file.name

            return send_file(tmp_file_path, as_attachment=True, download_name=new_document, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


    return render_template('templates.html')

def read_document(filename):
    doc = docx.Document(filename)
    completed_text = [] 

    for paragraph in doc.paragraphs:
        completed_text.append(re.findall('({.*?})', paragraph.text))

    for x in completed_text:
        if len(x) > 1:
            x = [y.split(',') for y in x]
    completed_text = [item for sublist in completed_text for item in sublist]
    
    return completed_text

if __name__ == "__main__":
    app.run(debug = True)